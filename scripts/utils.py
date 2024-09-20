from PyPDF2 import PdfReader
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
import uuid
from readability import Document


def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text or ""
    return text


def clean_whitespace(text):
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def get_text_from_webpage(url):
    response = requests.get(url)
    doc = Document(response.text)

    # Extract the main content
    main_content = doc.summary()

    # Parse the main content with BeautifulSoup to clean it up
    soup = BeautifulSoup(main_content, "html.parser")
    cleaned_text = clean_whitespace(soup.get_text(separator="\n"))

    return cleaned_text


def generate_id():
    return int(str(uuid.uuid4().int & (2**63 - 1))[:7])
